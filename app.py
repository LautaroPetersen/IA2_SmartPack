import streamlit as st
st.set_page_config(page_title="Redistribución con IA", layout="wide")

import pandas as pd
import toml
from io import BytesIO, StringIO
from docx import Document
import google.generativeai as genai

# -------------------- CONFIGURACIÓN --------------------

secrets = toml.load("secrets.toml")
api_key = secrets["gemini"]["api_key"]
genai.configure(api_key=api_key)

# -------------------- FUNCIÓN IA --------------------

def obtener_respuesta_gemini(packing_list, pedidos, instrucciones, observaciones):
    prompt = f"""
Eres un experto en logística y distribución inteligente de productos. Se te proporcionará un listado de cajas (Packing List) con su contenido original, y una tabla con los pedidos de diferentes clientes.

Tu tarea es:
1. Analizar ambos documentos.
2. Proponer la redistribución más eficiente de las cajas entre los clientes, intentando minimizar la división del contenido original si es posible y respetando observaciones e instrucciones específicas. 
3. Indicar si queda mercadería en stock, que no sera redistribuida	a ninguno de los clientes.
4. Generar una explicación clara de la lógica utilizada.
5. Generar una tabla de redistribución con las cajas y sus contenidos, y referencia a la caja original.

6. Generar el contenido de las etiquetas para cada nueva caja, separadas por "---" entre una y otra. Cada etiqueta debe incluir:
- Todo lo descripto en las instrucciones adicionales y observaciones
- Nombre del cliente
- Número de caja nueva
- Caja original de referencia
- Descripción del producto
- Talle/variante + cantidad


⚠️ Importante:
- La tabla debe ser completa, sin omisiones ni cortes.
- No resumir resultados, mostrar todas las filas aunque sean muchas.
- El contenido debe estar correctamente alineado con separadores `|` para que pueda ser procesado como tabla markdown.
- Separar las tres secciones con los encabezados:
  EXPLICACIÓN:
  TABLA DE REDISTRIBUCIÓN:
  ETIQUETAS:

Datos proporcionados:
Packing List:
{packing_list}

Pedidos de Clientes:
{pedidos}

Instrucciones adicionales:
{instrucciones}

Observaciones:
{observaciones}
"""
    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content(
        prompt,
        generation_config={
            "max_output_tokens": 8192,
            "temperature": 0.7,
            "top_p": 0.9,
            "top_k": 40
        }
    )
    return response.text

# -------------------- PROCESAR RESPUESTA --------------------

def procesar_respuesta_gemini(respuesta):
    explicacion, tabla, etiquetas = "", "", []
    partes = respuesta.split("ETIQUETAS:")
    if len(partes) == 2:
        etiquetas = [e.strip() for e in partes[1].split("---") if e.strip()]
    cuerpo = partes[0].split("TABLA DE REDISTRIBUCIÓN:")
    if len(cuerpo) == 2:
        explicacion = cuerpo[0].replace("EXPLICACIÓN:", "").strip()
        tabla = cuerpo[1].strip()
    return explicacion, tabla, etiquetas

# -------------------- FUNCIONES DE ARCHIVOS --------------------

def tabla_markdown_a_df(tabla_md):
    if "|" not in tabla_md:
        raise ValueError("La tabla generada por la IA no tiene formato markdown válido.")
    lineas = [l for l in tabla_md.strip().splitlines() if "|" in l and "---" not in l]
    if len(lineas) < 3:
        raise ValueError("La tabla es demasiado corta o mal estructurada.")
    tabla_limpia = "\n".join(lineas)
    try:
        df = pd.read_csv(StringIO(tabla_limpia), sep="|")
        df = df.iloc[:, 1:-1]
        df.columns = [c.strip() for c in df.columns]
        return df.dropna(how="all").reset_index(drop=True)
    except Exception as e:
        raise ValueError(f"No se pudo procesar la tabla: {e}")

def generar_excel(df):
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Redistribución")
    buffer.seek(0)
    return buffer

def generar_docx(etiquetas):
    doc = Document()
    for etiqueta in etiquetas:
        for linea in etiqueta.splitlines():
            doc.add_paragraph(linea)
        doc.add_paragraph("-" * 40)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -------------------- STREAMLIT UI --------------------

st.title("📦 Redistribución Inteligente de Productos con IA")

with st.expander("ℹ️ Cómo funciona"):
    st.markdown("""
Esta aplicación utiliza inteligencia artificial (Gemini) para redistribuir automáticamente el contenido de un pedido entre varios clientes, a partir de un packing list y sus respectivos pedidos.

**¿Qué hace?**
- Analiza el contenido original de las cajas
- Interpreta los pedidos de cada cliente
- Propone la redistribución más eficiente (intentando no dividir cajas innecesariamente)

**¿Qué devuelve?**
- Una explicación estratégica de la redistribución
- Un archivo Excel con el detalle completo de la redistribución
- Un archivo Word con las etiquetas para cada caja nueva

**¿Qué necesitás subir?**
1. Packing List original (.xlsx)
2. Tabla de pedidos por cliente (.xlsx)
3. (Opcional) Instrucciones para el formato de etiquetas
4. (Opcional) Observaciones especiales como límites por caja, prioridades, etc.
""")

packing_file = st.file_uploader("📥 Packing List (.xlsx)", type=["xlsx"])
pedidos_file = st.file_uploader("📥 Pedidos de Clientes (.xlsx)", type=["xlsx"])
instrucciones = st.text_area("📌 Instrucciones para etiquetas (opcional)")
observaciones = st.text_area("📝 Observaciones adicionales (opcional)")

if st.button("🚀 Generar Redistribución con IA"):
    if not packing_file or not pedidos_file:
        st.warning("Por favor cargá ambos archivos.")
    else:
        packing_df = pd.read_excel(packing_file)
        pedidos_df = pd.read_excel(pedidos_file)
        st.session_state["packing_list"] = packing_df.to_markdown(index=False)
        st.session_state["pedidos"] = pedidos_df.to_markdown(index=False)

        with st.spinner("🧠 Pensando con IA..."):
            respuesta = obtener_respuesta_gemini(
                st.session_state["packing_list"],
                st.session_state["pedidos"],
                instrucciones,
                observaciones
            )
            explicacion, tabla, etiquetas = procesar_respuesta_gemini(respuesta)
            st.session_state["explicacion"] = explicacion
            st.session_state["tabla_md"] = tabla
            st.session_state["etiquetas"] = etiquetas

            try:
                df_tabla = tabla_markdown_a_df(tabla)
                st.session_state["df_tabla"] = df_tabla
                st.success("✅ Redistribución generada correctamente.")
            except Exception as e:
                st.error(f"❌ Error al procesar la tabla: {e}")
                st.stop()

# Mostrar resultados y descargas
if "explicacion" in st.session_state:
    st.subheader("🧠 Explicación generada por la IA")
    st.markdown(st.session_state["explicacion"])

if "df_tabla" in st.session_state:
    excel_file = generar_excel(st.session_state["df_tabla"])
    st.download_button("📥 Descargar Excel de Redistribución", data=excel_file, file_name="redistribucion.xlsx")

if "etiquetas" in st.session_state:
    etiquetas_docx = generar_docx(st.session_state["etiquetas"])
    st.download_button("📄 Descargar Etiquetas (Word)", data=etiquetas_docx, file_name="etiquetas.docx")
